import pandas as pd
import os
import numpy as np
import sys
sys.path.append('data_analysis/helper')
from clean_data import get_prolific_ids_with_few_user_turns
from docx import Document

def extract_conversations(df, output_dir='data_analysis/conversation/conversations/', min_user_turns=3):
    """
    Extract conversations for each prolific ID and save them to files.
    
    Args:
        df (pd.DataFrame): DataFrame with conversation data
        output_dir (str): Directory to save conversation files
        min_user_turns (int): Minimum number of user turns required (default: 3)
    """
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Get prolific IDs with few user turns to exclude them
    prolific_ids_to_exclude = get_prolific_ids_with_few_user_turns(min_user_turns)
    print(f"Will exclude {len(prolific_ids_to_exclude)} prolific IDs with fewer than {min_user_turns} user turns")
    
    # Debug: Check data structure
    print(f"DataFrame shape: {df.shape}")
    print(f"Columns: {df.columns.tolist()}")
    print(f"Event types: {df['event_type'].value_counts()}")
    print(f"Unique sessions: {df['session_id'].nunique()}")
    print(f"Unique prolific IDs: {df['prolific_id'].nunique()}")
    
    # Filter for message events only and sort by session and timestamp
    messages_df = df[df['event_type'] == 'message'].copy()
    print(f"Message events: {len(messages_df)}")
    
    if len(messages_df) == 0:
        print("No message events found!")
        return {}
    
    messages_df = messages_df.sort_values(['session_id', 'timestamp'])
    
    # Debug: Check message data
    print(f"Message roles: {messages_df['role'].value_counts()}")
    print(f"Messages with content: {messages_df['content'].notna().sum()}")
    
    # Group by session_id to process each conversation
    conversations = {}
    
    for session_id, session_data in messages_df.groupby('session_id'):
        print(f"\nProcessing session: {session_id}")
        
        # Get session metadata from the first row
        prolific_id = session_data['prolific_id'].iloc[0]
        dialogue_mode = session_data['dialogue_mode'].iloc[0]
        origin_url = session_data['origin_url'].iloc[0]
        article = session_data['article'].iloc[0]
        
        print(f"  Prolific ID: {prolific_id}")
        print(f"  Dialogue Mode: {dialogue_mode}")
        print(f"  Messages in session: {len(session_data)}")
        
        if pd.isna(prolific_id):
            print(f"  Skipping session {session_id} - no prolific_id")
            continue
            
        # Build conversation text
        conversation_lines = []
        conversation_lines.append(f"Session ID: {session_id}")
        conversation_lines.append(f"Dialogue Mode: {dialogue_mode}")
        conversation_lines.append(f"Prolific ID: {prolific_id}")
        conversation_lines.append(f"URL: {origin_url}")
        # conversation_lines.append(f"Article: {article}")
        conversation_lines.append("-" * 50)
        conversation_lines.append("")
        
        message_count = 0
        for _, row in session_data.iterrows():
            role = row['role']
            content = row['content']
            
            if pd.notna(content) and pd.notna(role):
                conversation_lines.append(f"{role}: {content}")
                conversation_lines.append("")
                message_count += 1
        
        print(f"  Valid messages: {message_count}")
        
        # Join all lines
        conversation_text = "\n".join(conversation_lines)
        
        # Store by dialogue_mode and prolific_id
        if dialogue_mode not in conversations:
            conversations[dialogue_mode] = {}
        if prolific_id not in conversations[dialogue_mode]:
            conversations[dialogue_mode][prolific_id] = []
        conversations[dialogue_mode][prolific_id].append(conversation_text)
    
    # Save conversations to files
    saved_count = 0
    excluded_count = 0
    for dialogue_mode, prolific_conversations in conversations.items():
        # Create directory for this dialogue mode
        mode_dir = os.path.join(output_dir, dialogue_mode)
        os.makedirs(mode_dir, exist_ok=True)
        
        for prolific_id, conversation_list in prolific_conversations.items():
            # Skip prolific IDs with few user turns
            if prolific_id in prolific_ids_to_exclude:
                excluded_count += 1
                print(f"Skipping prolific_id {prolific_id} (mode: {dialogue_mode}) - too few user turns")
                continue
                
            filename = f"{mode_dir}/conversation_{prolific_id}.docx"
            
            # If multiple sessions for same prolific_id, combine them
            if len(conversation_list) > 1:
                combined_text = "\n\n" + "="*80 + "\n\n".join(conversation_list)
            else:
                combined_text = conversation_list[0]
            
            # Create a new Word document
            doc = Document()
            
            # Split the text into lines and add to document
            lines = combined_text.split('\n')
            for line in lines:
                if line.strip():  # Only add non-empty lines
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph()  # Add empty paragraph for spacing
            
            # Save the document
            doc.save(filename)
            
            print(f"Saved conversation for prolific_id {prolific_id} (mode: {dialogue_mode}) to {filename}")
            saved_count += 1
    
    print(f"\nTotal conversations saved: {saved_count}")
    print(f"Total conversations excluded: {excluded_count}")
    return conversations

def get_conversation_stats(df):
    """
    Get statistics about conversations.
    """
    messages_df = df[df['event_type'] == 'message'].copy()
    
    if len(messages_df) == 0:
        print("No message events found for statistics!")
        return
    
    # Count messages per session
    messages_per_session = messages_df.groupby('session_id').size()
    
    # Count sessions per prolific_id
    sessions_per_prolific = messages_df.groupby('prolific_id')['session_id'].nunique()
    
    print("Conversation Statistics:")
    print(f"Total sessions: {len(messages_per_session)}")
    print(f"Total prolific IDs: {len(sessions_per_prolific)}")
    print(f"Average messages per session: {messages_per_session.mean():.2f}")
    print(f"Median messages per session: {messages_per_session.median():.2f}")
    print(f"Min messages per session: {messages_per_session.min()}")
    print(f"Max messages per session: {messages_per_session.max()}")
    print(f"Average sessions per prolific ID: {sessions_per_prolific.mean():.2f}")
    
    # Show distribution
    print("\nMessages per session distribution:")
    print(messages_per_session.value_counts().sort_index())

# Example usage with your real data
if __name__ == "__main__":
    # Load your DynamoDB data
    import boto3
    import os
    from datetime import datetime
    from dotenv import load_dotenv

    # Load environment variables from .env file (if it exists)
    load_dotenv()

    # Initialize DynamoDB client
    aws_region = os.getenv('AWS_REGION', 'eu-north-1')
    endpoint_url = os.getenv('AWS_ENDPOINT_URL')
    DYNAMODB_TABLE = os.getenv('DYNAMODB_TABLE', 'apollolytics_dialogues')

    if endpoint_url:
        dynamodb = boto3.resource('dynamodb', region_name=aws_region, endpoint_url=endpoint_url)
    else:
        dynamodb = boto3.resource('dynamodb', region_name=aws_region)

    # Get the table
    table = dynamodb.Table(DYNAMODB_TABLE)

    # Scan the table to get all items
    response = table.scan()
    items = response['Items']

    # Handle pagination if there are more results
    while 'LastEvaluatedKey' in response:
        response = table.scan(ExclusiveStartKey=response['LastEvaluatedKey'])
        items.extend(response['Items'])

    # Convert to DataFrame
    df = pd.DataFrame(items)

    # Convert timestamp to datetime for better analysis
    df['timestamp'] = pd.to_numeric(df['timestamp'], errors='coerce')
    df['datetime'] = pd.to_datetime(df['timestamp'], unit='s')

    # Sort by session_id and timestamp
    df = df.sort_values(['session_id', 'timestamp'])

    # Filter for valid prolific IDs (24 characters)
    session_ids = df[df.prolific_id.str.len() == 24].session_id
    df = df[df.session_id.isin(session_ids)]
    
    # Remove columns with only NaN values
    df = df.dropna(axis=1, how='all')
    print(df.columns)
    # Forward fill session-specific columns within each session
    session_columns = ['dialogue_mode', 'origin_url', 'article', 'prolific_id']
    df[session_columns] = df.groupby('session_id')[session_columns].ffill()

    print("Data loaded and processed!")
    print(f"DataFrame shape: {df.shape}")
    
    # Get conversation statistics
    get_conversation_stats(df)
    
    # Extract and save conversations
    conversations = extract_conversations(df, 'data_analysis/conversation/conversations/', min_user_turns=3) 